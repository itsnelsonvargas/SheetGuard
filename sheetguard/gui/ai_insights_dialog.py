"""Dialog to display AI Insights."""

from PySide6.QtWidgets import (
    QDialog,
    QFileDialog,
    QHBoxLayout,
    QMessageBox,
    QPushButton,
    QTextEdit,
    QVBoxLayout,
)

class AIInsightsDialog(QDialog):
    """Dialog to render markdown insights from the AI."""
    
    def __init__(self, insights_markdown: str, parent=None) -> None:
        super().__init__(parent)
        self.setWindowTitle("AI Data Insights")
        self.setMinimumSize(700, 500)
        self.setObjectName("aiInsightsDialog")
        
        self.insights_markdown = insights_markdown
        self._build_ui()

    def _build_ui(self) -> None:
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        
        self.text_edit = QTextEdit()
        self.text_edit.setReadOnly(True)
        # QTextEdit has limited markdown support, but sufficient for basic bold/lists
        self.text_edit.setMarkdown(self.insights_markdown)
        self.text_edit.setObjectName("aiInsightsText")
        # Ensure it has a white/dark background based on theme instead of transparent
        self.text_edit.setStyleSheet("font-size: 14px; padding: 10px;")
        layout.addWidget(self.text_edit)
        
        btn_row = QHBoxLayout()
        
        btn_export = QPushButton("Export to File...")
        btn_export.setObjectName("secondary")
        btn_export.clicked.connect(self._export_review)
        btn_row.addWidget(btn_export)
        
        btn_row.addStretch()
        
        btn_close = QPushButton("Close")
        btn_close.setObjectName("success")
        btn_close.clicked.connect(self.accept)
        btn_row.addWidget(btn_close)
        
        layout.addLayout(btn_row)

    def _export_review(self) -> None:
        filters = (
            "Markdown Files (*.md);;"
            "PDF Files (*.pdf);;"
            "Word Documents (*.docx);;"
            "Text Files (*.txt);;"
            "All Files (*)"
        )
        path, _ = QFileDialog.getSaveFileName(
            self, "Export AI Review", "ai_review_insights.md", filters
        )
        if not path:
            return
            
        try:
            if path.endswith(".pdf"):
                from PySide6.QtPrintSupport import QPrinter
                printer = QPrinter(QPrinter.PrinterMode.HighResolution)
                printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
                printer.setOutputFileName(path)
                self.text_edit.document().print_(printer)
            elif path.endswith(".docx"):
                from docx import Document
                doc = Document()
                doc.add_heading('AI Data Insights', 0)
                for line in self.text_edit.toPlainText().split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
                doc.save(path)
            else:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(self.insights_markdown)
            QMessageBox.information(self, "Export Successful", f"Saved review to:\n{path}")
        except Exception as exc:
            QMessageBox.critical(self, "Export Error", f"Could not save file:\n{exc}")
